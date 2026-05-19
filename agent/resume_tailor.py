import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import copy

logger = logging.getLogger(__name__)

class ResumeTailor:
    def __init__(self, user):
        self.user = user
        os.makedirs("resumes\\tailored", exist_ok=True)

    def tailor(self, job):
        return self._build_docx(job)

    def _build_docx(self, job):
        resume_path = self.user.resume_path

        if resume_path and os.path.exists(resume_path) and resume_path.endswith(".docx"):
            doc = Document(resume_path)
        elif resume_path and os.path.exists(resume_path) and resume_path.endswith(".pdf"):
            doc = self._create_basic_doc()
        else:
            doc = self._create_basic_doc()

        if job.description:
            p = doc.add_paragraph()
            run = p.add_run(job.description)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(1)

        safe = lambda s: "".join(c for c in s if c.isalnum() or c in " _-")[:25].strip()
        filename = f"{safe(job.company)}_{safe(job.title)}_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.docx".replace(" ", "_")
        path = os.path.join("resumes\\tailored", filename)
        doc.save(path)
        logger.info(f"Resume saved: {path}")
        return path

    def _create_basic_doc(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for section in doc.sections:
            section.top_margin = Pt(36)
            section.bottom_margin = Pt(36)
            section.left_margin = Pt(54)
            section.right_margin = Pt(54)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.user.full_name)
        run.bold = True
        run.font.size = Pt(20)
        contact = " | ".join(filter(None, [self.user.email, self.user.phone, self.user.linkedin_url]))
        p = doc.add_paragraph(contact)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Target Roles: {self.user.target_roles}")
        return doc

    def _read_resume(self):
        path = self.user.resume_path
        if not path or not os.path.exists(path):
            return f"Name: {self.user.full_name}, Target: {self.user.target_roles}"
        try:
            if path.endswith(".pdf"):
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            elif path.endswith(".docx"):
                d = Document(path)
                return "\n".join(p.text for p in d.paragraphs)
        except Exception as e:
            logger.error(f"Resume read error: {e}")
        return ""